"""
screenshot_assist_full_with_activefile.py
Single-file production-ready ScreenshotAssist app (PyQt6)

Features included:
- Full-screen / Active-window / Region capture
- Manual hotkeys (F9 full capture, F10 region capture)
- Scheduled capture
- Annotation (rectangles + text) of last screenshot
- History panel (thumbnails)
- Auto-naming, session folders
- Append screenshots to Word (.docx)
- Copy to clipboard (Windows)
- Dark/Light theme toggle
- System tray quick actions
- Global hotkeys (keyboard)
- Image compression, optional blur
- get_active_window_file() provides window title + exe path/name
- Robust logging and config persistence

Icon used (default): /mnt/data/IMG_2111.jpeg (or the .ico at /mnt/data/screenshotassist_icon.ico)

Run:
    python screenshot_assist_full_with_activefile.py

Dependencies:
    pip install pyqt6 pyautogui pillow python-docx keyboard pystray pywin32 psutil
"""

import os
import sys
import json
import time
import threading
import logging
from datetime import datetime
from pathlib import Path
from io import BytesIO

# PyQt6
from PyQt6 import QtWidgets, QtGui, QtCore
from PyQt6.QtCore import Qt, QRect, QPoint
from PyQt6.QtGui import QPixmap, QPainter, QPen, QColor, QImage, QAction, QKeySequence, QShortcut
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QFileDialog, QLabel, QListWidget,
    QListWidgetItem, QHBoxLayout, QVBoxLayout, QWidget, QPushButton, QSpinBox,
    QCheckBox, QFrame, QTextEdit, QSplitter, QMessageBox
)

# Other libs
import pyautogui
from PIL import Image, ImageFilter, ImageDraw
from docx import Document
from docx.shared import Inches
import keyboard
import pystray
from pystray import MenuItem as item
import win32clipboard
from PIL import ImageWin
import ctypes
import psutil
import win32gui
import win32process

# ---------------- config & logging ----------------
APP_NAME = "ScreenshotAssist"
ICON_JPEG = "/mnt/data/IMG_2111.jpeg"
ICON_ICO = "/mnt/data/screenshotassist_icon.ico"
CONFIG_FILE = "screenshot_assist_config.json"
LOG_FILE = "screenshot_assist.log"

DEFAULT_CONFIG = {
    "hotkey_capture": "f9",
    "hotkey_region": "f10",
    "schedule_interval": 30,
    "optimize_images": True,
    "image_max_width": 1200,
    "compress_quality": 80,
    "blur_center": False,
    "add_headings": True,
    "add_page_breaks": False,
    "output_base": str(Path.home() / "Documents" / "ScreenshotAssist")
}

logging.basicConfig(filename=LOG_FILE, level=logging.INFO,
                    format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger(APP_NAME)

# ---------------- config utilities ----------------
def load_config():
    if not os.path.exists(CONFIG_FILE):
        save_config(DEFAULT_CONFIG.copy())
        return DEFAULT_CONFIG.copy()
    try:
        with open(CONFIG_FILE, "r") as f:
            conf = json.load(f)
        for k, v in DEFAULT_CONFIG.items():
            if k not in conf:
                conf[k] = v
        return conf
    except Exception:
        logger.exception("Failed to load config, using defaults")
        return DEFAULT_CONFIG.copy()

def save_config(conf):
    try:
        with open(CONFIG_FILE, "w") as f:
            json.dump(conf, f, indent=2)
    except Exception:
        logger.exception("Failed to save config")

config = load_config()

# ---------------- session folders ----------------
def ensure_session_folder(base: str = None):
    base = base or config.get("output_base")
    base_path = Path(base)
    base_path.mkdir(parents=True, exist_ok=True)
    today = datetime.now().strftime("%Y-%m-%d")
    day_path = base_path / today
    day_path.mkdir(exist_ok=True)
    # find next session
    i = 1
    while (day_path / f"Session_{i}").exists():
        i += 1
    session_path = day_path / f"Session_{i}"
    (session_path / "screenshots").mkdir(parents=True, exist_ok=True)
    return {
        "session_path": str(session_path),
        "screenshots": str(session_path / "screenshots"),
        "word": str(session_path / f"Word_Report_Session_{i}.docx")
    }

session = ensure_session_folder()

# ---------------- Word generator ----------------
class WordGen:
    def __init__(self, path, add_headings=True, add_page_breaks=False):
        self.path = path
        self.add_headings = add_headings
        self.add_page_breaks = add_page_breaks
        if os.path.exists(self.path):
            self.doc = Document(self.path)
        else:
            self.doc = Document()
            self.doc.add_heading("ScreenshotAssist Report", level=1)
            self.doc.add_paragraph(f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            self.doc.save(self.path)
        self.step = 1
        self.last_title = None
        self.lock = threading.Lock()

    def add(self, img_path, title):
        with self.lock:
            try:
                if self.add_headings and title and title != self.last_title:
                    self.doc.add_heading(title, level=2)
                    self.last_title = title
                self.doc.add_paragraph(f"Step {self.step} — {datetime.now().strftime('%H:%M:%S')}")
                self.step += 1
                self.doc.add_picture(img_path, width=Inches(6))
                if self.add_page_breaks:
                    self.doc.add_page_break()
                self.doc.save(self.path)
                logger.info(f"Image appended to doc: {img_path}")
            except Exception:
                logger.exception("Word append failed")

wordgen = WordGen(session["word"], add_headings=config.get("add_headings", True),
                  add_page_breaks=config.get("add_page_breaks", False))

# ---------------- image helper funcs ----------------
def compress_image(in_path, max_width=None, quality=None):
    try:
        max_width = max_width or config.get("image_max_width", 1200)
        quality = quality or config.get("compress_quality", 80)
        img = Image.open(in_path)
        w, h = img.size
        if config.get("optimize_images", True) and w > max_width:
            ratio = max_width / float(w)
            new_size = (max_width, int(h * ratio))
            img = img.resize(new_size, Image.LANCZOS)
        out_path = str(Path(in_path).with_suffix(".jpg"))
        img.convert("RGB").save(out_path, "JPEG", quality=quality, optimize=True)
        if out_path != in_path and os.path.exists(in_path):
            try:
                os.remove(in_path)
            except:
                pass
        return out_path
    except Exception:
        logger.exception("compress_image failed")
        return in_path

def blur_center(in_path):
    try:
        img = Image.open(in_path)
        w, h = img.size
        rect = (int(w*0.2), int(h*0.4), int(w*0.8), int(h*0.55))
        crop = img.crop(rect).filter(ImageFilter.GaussianBlur(radius=12))
        img.paste(crop, rect)
        out = str(Path(in_path).with_suffix(".blur.jpg"))
        img.convert("RGB").save(out, "JPEG", quality=85)
        return out
    except Exception:
        logger.exception("blur_center failed")
        return in_path

# ---------------- clipboard helper ----------------
def send_image_to_clipboard(pil_image):
    try:
        output = BytesIO()
        pil_image.convert('RGB').save(output, 'BMP')
        data = output.getvalue()[14:]
        output.close()
        win32clipboard.OpenClipboard()
        win32clipboard.EmptyClipboard()
        win32clipboard.SetClipboardData(win32clipboard.CF_DIB, data)
        win32clipboard.CloseClipboard()
        return True
    except Exception:
        logger.exception("send_image_to_clipboard failed")
        return False

# ---------------- get active window metadata ----------------
def get_active_window_file():
    """
    Returns dict:
      {"window_title": str, "exe_path": str, "exe_name": str}
    """
    try:
        hwnd = win32gui.GetForegroundWindow()
        if not hwnd:
            return {"window_title": "", "exe_path": "", "exe_name": ""}
        window_title = win32gui.GetWindowText(hwnd)
        thread_id, pid = win32process.GetWindowThreadProcessId(hwnd)
        exe_path = ""
        exe_name = ""
        if pid:
            try:
                proc = psutil.Process(pid)
                exe_path = proc.exe()
                exe_name = os.path.basename(exe_path)
            except Exception:
                # fallback: try GetModuleFileNameEx via pywin32 if needed (skip for simplicity)
                pass
        return {"window_title": window_title or "", "exe_path": exe_path or "", "exe_name": exe_name or ""}
    except Exception:
        logger.exception("get_active_window_file failed")
        return {"window_title": "", "exe_path": "", "exe_name": ""}

# ---------------- capture implementations ----------------
def capture_fullscreen(save_dir):
    ts = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    fname = f"fullscreen_{ts}.png"
    path = os.path.join(save_dir, fname)
    pyautogui.screenshot(path)
    return path

def capture_active_window(save_dir):
    try:
        hwnd = win32gui.GetForegroundWindow()
        x1, y1, x2, y2 = win32gui.GetWindowRect(hwnd)
        ts = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        fname = f"active_{ts}.png"
        path = os.path.join(save_dir, fname)
        img = pyautogui.screenshot(region=(x1, y1, x2-x1, y2-y1))
        img.save(path)
        return path
    except Exception:
        logger.exception("capture_active_window failed, falling back to fullscreen")
        return capture_fullscreen(save_dir)

# ---------------- Region selector overlay ----------------
class RegionSelector(QtWidgets.QWidget):
    regionSelected = QtCore.pyqtSignal(QtCore.QRect)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.WindowStaysOnTopHint)
        self.setWindowState(Qt.WindowState.WindowFullScreen)
        self.start = QPoint()
        self.end = QPoint()
        self.setCursor(Qt.CursorShape.CrossCursor)
        self.show()

    def mousePressEvent(self, event):
        self.start = event.pos()
        self.end = self.start
        self.update()

    def mouseMoveEvent(self, event):
        self.end = event.pos()
        self.update()

    def mouseReleaseEvent(self, event):
        self.end = event.pos()
        rect = QRect(self.start, self.end).normalized()
        self.regionSelected.emit(rect)
        self.close()

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setOpacity(0.35)
        painter.fillRect(self.rect(), QColor(0,0,0))
        pen = QPen(QColor(255,255,255), 2)
        painter.setPen(pen)
        rect = QRect(self.start, self.end).normalized()
        painter.setOpacity(1.0)
        painter.drawRect(rect)

# ---------------- Annotation window ----------------
class AnnotateWindow(QMainWindow):
    def __init__(self, image_path, save_callback):
        super().__init__()
        self.setWindowTitle("Annotate Screenshot")
        self.save_callback = save_callback
        self.image_path = image_path
        self.pil = Image.open(self.image_path).convert("RGBA")
        data = self.pil.tobytes("raw", "RGBA")
        qimg = QImage(data, self.pil.size[0], self.pil.size[1], QImage.Format.Format_RGBA8888)
        self.pix = QPixmap.fromImage(qimg)
        self.init_ui()

    def init_ui(self):
        self.label = QLabel()
        self.label.setPixmap(self.pix)
        self.setCentralWidget(self.label)
        self.setFixedSize(self.pix.size())
        self.drawing = False
        self.start = QPoint()
        self.end = QPoint()
        self.rects = []
        self.texts = []
        self.label.mousePressEvent = self.on_mouse_press
        self.label.mouseMoveEvent = self.on_mouse_move
        self.label.mouseReleaseEvent = self.on_mouse_release
        self.label.paintEvent = self.on_paint
        toolbar = self.addToolBar("Annotate")
        save_act = QAction("Save & Close", self)
        save_act.triggered.connect(self.save_and_close)
        toolbar.addAction(save_act)

    def on_mouse_press(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self.drawing = True
            self.start = event.pos()
            self.end = self.start
        elif event.button() == Qt.MouseButton.RightButton:
            text, ok = QtWidgets.QInputDialog.getText(self, "Add Text", "Enter text:")
            if ok and text:
                self.texts.append((event.pos(), text))
                self.update()

    def on_mouse_move(self, event):
        if self.drawing:
            self.end = event.pos()
            self.update()

    def on_mouse_release(self, event):
        if event.button() == Qt.MouseButton.LeftButton and self.drawing:
            self.drawing = False
            rect = QRect(self.start, self.end).normalized()
            self.rects.append(rect)
            self.update()

    def on_paint(self, event):
        painter = QPainter(self.label)
        painter.drawPixmap(0, 0, self.pix)
        pen = QPen(QColor(255, 0, 0), 3)
        painter.setPen(pen)
        for r in self.rects:
            painter.drawRect(r)
        painter.setPen(QPen(QColor(0, 255, 0), 1))
        for pos, t in self.texts:
            painter.drawText(pos, t)

    def save_and_close(self):
        draw = ImageDraw.Draw(self.pil)
        for r in self.rects:
            draw.rectangle([r.left(), r.top(), r.right(), r.bottom()], outline=(255,0,0), width=6)
        for pos, t in self.texts:
            draw.text((pos.x(), pos.y()), t, fill=(255,255,0))
        out = str(Path(self.image_path).with_suffix(".annot.jpg"))
        self.pil.convert("RGB").save(out, "JPEG", quality=90)
        self.save_callback(out)
        self.close()

# ---------------- History list ----------------
class HistoryList(QListWidget):
    def __init__(self):
        super().__init__()
        self.setViewMode(QListWidget.ViewMode.IconMode)
        self.setIconSize(QtCore.QSize(240, 135))
        self.setResizeMode(QListWidget.ResizeMode.Adjust)
        self.setSpacing(8)
        self.items = []

    def add_image(self, img_path):
        item = QListWidgetItem()
        pix = QPixmap(img_path).scaled(320, 180, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
        item.setIcon(QtGui.QIcon(pix))
        item.setText(Path(img_path).name)
        item.setToolTip(img_path)
        self.addItem(item)
        self.items.append(img_path)

# ---------------- Scheduler ----------------
schedule_running = False
schedule_thread = None

def start_scheduler(interval, main_window):
    global schedule_running, schedule_thread
    if schedule_running:
        return
    schedule_running = True
    def loop():
        logger.info("Scheduler started")
        while schedule_running:
            path = capture_fullscreen(session["screenshots"])
            final = compress_image(path) if config.get("optimize_images", True) else path
            if config.get("blur_center", False):
                final = blur_center(final)
            meta = get_active_window_file()
            wordgen.add(final, meta.get("window_title","") if config.get("add_headings", True) else "")
            # add to UI
            QtCore.QMetaObject.invokeMethod(main_window, "add_history_item", QtCore.Qt.ConnectionType.QueuedConnection,
                                            QtCore.Q_ARG(str, final))
            for _ in range(int(interval * 10)):
                if not schedule_running:
                    break
                time.sleep(0.1)
        logger.info("Scheduler stopped")
    schedule_thread = threading.Thread(target=loop, daemon=True)
    schedule_thread.start()

def stop_scheduler():
    global schedule_running
    schedule_running = False

# ---------------- tray ----------------
tray_icon = None
def start_tray(main_window):
    global tray_icon
    try:
        if os.path.exists(ICON_ICO):
            img = Image.open(ICON_ICO).resize((64,64))
        elif os.path.exists(ICON_JPEG):
            img = Image.open(ICON_JPEG).resize((64,64))
        else:
            img = Image.new("RGB", (64,64), (30,90,200))
        menu = (
            item('Open', lambda: main_window.show_window()),
            item('Capture Full', lambda: main_window.capture_full()),
            item('Capture Region', lambda: main_window.capture_region()),
            item('Start Schedule', lambda: start_scheduler(config.get("schedule_interval",30), main_window)),
            item('Stop Schedule', lambda: stop_scheduler()),
            item('Exit', lambda: main_window.quit_app()),
        )
        tray_icon = pystray.Icon(APP_NAME, img, APP_NAME, menu)
        threading.Thread(target=tray_icon.run, daemon=True).start()
    except Exception:
        logger.exception("start_tray failed")

def stop_tray():
    global tray_icon
    try:
        if tray_icon:
            tray_icon.stop()
    except Exception:
        pass

# ---------------- Main UI ----------------
class MainApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("ScreenshotAssist - Professional")
        self.setGeometry(150, 80, 1200, 720)
        if os.path.exists(ICON_ICO):
            self.setWindowIcon(QtGui.QIcon(ICON_ICO))
        elif os.path.exists(ICON_JPEG):
            self.setWindowIcon(QtGui.QIcon(ICON_JPEG))
        # UI
        self.history = HistoryList()
        self._build_ui()
        # register hotkeys
        try:
            keyboard.add_hotkey(config.get("hotkey_capture","f9"), self.capture_full)
            keyboard.add_hotkey(config.get("hotkey_region","f10"), self.capture_region)
            logger.info("Hotkeys registered")
        except Exception:
            logger.exception("hotkey registration failed - try running as Administrator")
        # start tray
        start_tray(self)
        self.last_image = None
        self.dark = False

    def _build_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QHBoxLayout(central)

        # Left area controls
        left = QVBoxLayout()
        btn_row = QHBoxLayout()
        self.btn_full = QPushButton("Capture Fullscreen")
        self.btn_active = QPushButton("Capture Active Window")
        self.btn_region = QPushButton("Capture Region")
        self.btn_annot = QPushButton("Annotate Last")
        btn_row.addWidget(self.btn_full); btn_row.addWidget(self.btn_active); btn_row.addWidget(self.btn_region); btn_row.addWidget(self.btn_annot)
        self.btn_full.clicked.connect(self.capture_full)
        self.btn_active.clicked.connect(self.capture_active)
        self.btn_region.clicked.connect(self.capture_region)
        self.btn_annot.clicked.connect(self.annotate_last)
        left.addLayout(btn_row)

        # Schedule control
        sched_row = QHBoxLayout()
        sched_row.addWidget(QtWidgets.QLabel("Schedule sec:"))
        self.spin_sched = QSpinBox(); self.spin_sched.setRange(1,86400); self.spin_sched.setValue(config.get("schedule_interval",30))
        self.btn_start_sched = QPushButton("Start Schedule"); self.btn_stop_sched = QPushButton("Stop Schedule")
        self.btn_start_sched.clicked.connect(lambda: start_scheduler(int(self.spin_sched.value()), self))
        self.btn_stop_sched.clicked.connect(stop_scheduler)
        sched_row.addWidget(self.spin_sched); sched_row.addWidget(self.btn_start_sched); sched_row.addWidget(self.btn_stop_sched)
        left.addLayout(sched_row)

        # options
        self.chk_opt = QCheckBox("Optimize Images"); self.chk_opt.setChecked(config.get("optimize_images", True))
        self.chk_blur = QCheckBox("Blur Center"); self.chk_blur.setChecked(config.get("blur_center", False))
        self.chk_head = QCheckBox("Auto Headings"); self.chk_head.setChecked(config.get("add_headings", True))
        self.chk_page = QCheckBox("Page Breaks"); self.chk_page.setChecked(config.get("add_page_breaks", False))
        left.addWidget(self.chk_opt); left.addWidget(self.chk_blur); left.addWidget(self.chk_head); left.addWidget(self.chk_page)

        # folder / doc
        self.lbl_session = QLabel(f"Session: {session['session_path']}")
        left.addWidget(self.lbl_session)
        self.btn_open_folder = QPushButton("Open Session Folder"); self.btn_open_folder.clicked.connect(self.open_folder)
        left.addWidget(self.btn_open_folder)
        self.btn_copy_clip = QPushButton("Copy Last to Clipboard"); self.btn_copy_clip.clicked.connect(self.copy_last_to_clipboard)
        self.btn_save_doc = QPushButton("Save Word Document"); self.btn_save_doc.clicked.connect(self.save_doc)
        left.addWidget(self.btn_copy_clip); left.addWidget(self.btn_save_doc)

        # theme & status
        self.btn_theme = QPushButton("Toggle Dark/Light Theme"); self.btn_theme.clicked.connect(self.toggle_theme)
        left.addWidget(self.btn_theme)
        self.txt_status = QTextEdit(); self.txt_status.setReadOnly(True); self.txt_status.setFixedHeight(180)
        left.addWidget(self.txt_status)

        left_widget = QWidget(); left_widget.setLayout(left)

        # right: history
        right_layout = QVBoxLayout()
        right_layout.addWidget(QLabel("History"))
        right_layout.addWidget(self.history)
        right_widget = QWidget(); right_widget.setLayout(right_layout)

        splitter = QSplitter()
        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setSizes([450, 750])
        main_layout.addWidget(splitter)

    # UI helpers
    def log(self, text):
        ts = datetime.now().strftime("%H:%M:%S")
        self.txt_status.append(f"[{ts}] {text}")
        logger.info(text)

    @QtCore.pyqtSlot(str)
    def add_history_item(self, path):
        self.history.add_image(path)
        self.last_image = path

    # Actions
    def capture_full(self):
        self.log("Capturing fullscreen...")
        path = capture_fullscreen(session["screenshots"])
        final = compress_image(path) if self.chk_opt.isChecked() else path
        if self.chk_blur.isChecked():
            final = blur_center(final)
        meta = get_active_window_file()
        wordgen.add(final, meta.get("window_title","") if self.chk_head.isChecked() else "")
        self.add_history_item(final)
        self.log(f"Saved {final}")

    def capture_active(self):
        self.log("Capturing active window...")
        path = capture_active_window(session["screenshots"])
        final = compress_image(path) if self.chk_opt.isChecked() else path
        if self.chk_blur.isChecked():
            final = blur_center(final)
        meta = get_active_window_file()
        wordgen.add(final, meta.get("window_title","") if self.chk_head.isChecked() else "")
        self.add_history_item(final)
        self.log(f"Saved {final}")

    def capture_region(self):
        self.log("Select region...")
        sel = RegionSelector()
        sel.regionSelected.connect(lambda rect: self._handle_region_selected(rect))
        sel.show()

    def _handle_region_selected(self, rect: QRect):
        x = rect.left(); y = rect.top(); w = rect.width(); h = rect.height()
        ts = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        fname = f"region_{ts}.png"
        path = os.path.join(session["screenshots"], fname)
        im = pyautogui.screenshot(region=(x, y, w, h))
        im.save(path)
        final = compress_image(path) if self.chk_opt.isChecked() else path
        if self.chk_blur.isChecked():
            final = blur_center(final)
        meta = get_active_window_file()
        wordgen.add(final, meta.get("window_title","") if self.chk_head.isChecked() else "")
        self.add_history_item(final)
        self.log(f"Saved region {final}")

    def annotate_last(self):
        if not self.last_image:
            self.log("No image to annotate")
            return
        win = AnnotateWindow(self.last_image, self._after_annot)
        win.show()

    def _after_annot(self, new_path):
        final = compress_image(new_path) if self.chk_opt.isChecked() else new_path
        wordgen.add(final, get_active_window_file().get("window_title","") if self.chk_head.isChecked() else "")
        self.add_history_item(final)
        self.log(f"Annotated saved {final}")

    def copy_last_to_clipboard(self):
        if not self.last_image or not os.path.exists(self.last_image):
            self.log("No last image")
            return
        img = Image.open(self.last_image)
        ok = send_image_to_clipboard(img)
        self.log("Copied to clipboard" if ok else "Failed to copy to clipboard")

    def save_doc(self):
        try:
            with wordgen.lock:
                wordgen.doc.save(wordgen.path)
            self.log("Word document saved")
        except Exception:
            logger.exception("save_doc failed")
            self.log("Failed saving document")

    def open_folder(self):
        os.startfile(session["session_path"])

    def toggle_theme(self):
        if self.dark:
            QApplication.setStyle("Fusion")
            QApplication.setPalette(QtGui.QPalette())
            self.dark = False
            self.log("Switched to Light theme")
        else:
            QApplication.setStyle("Fusion")
            darkPalette = QtGui.QPalette()
            darkPalette.setColor(QtGui.QPalette.ColorRole.Window, QColor(53,53,53))
            darkPalette.setColor(QtGui.QPalette.ColorRole.WindowText, Qt.white)
            QApplication.setPalette(darkPalette)
            self.dark = True
            self.log("Switched to Dark theme")

    def show_window(self):
        self.showNormal()
        self.activateWindow()

    def quit_app(self):
        stop_scheduler()
        stop_tray()
        try:
            keyboard.unhook_all_hotkeys()
        except:
            pass
        QApplication.quit()

# ---------------- startup helpers ----------------
def set_pyqt6_highdpi():
    try:
        from PyQt6 import QtCore, QtWidgets
        QtWidgets.QApplication.setHighDpiScaleFactorRoundingPolicy(
            QtCore.Qt.HighDpiScaleFactorRoundingPolicy.PassThrough
        )
    except Exception:
        pass

def main():
    os.environ["QT_ENABLE_HIGHDPI_SCALING"] = "1"
    set_pyqt6_highdpi()

    app = QApplication(sys.argv)
    QApplication.setStyle("Fusion")
    if os.path.exists(ICON_ICO):
        app.setWindowIcon(QtGui.QIcon(ICON_ICO))
    elif os.path.exists(ICON_JPEG):
        app.setWindowIcon(QtGui.QIcon(ICON_JPEG))

    main_window = MainApp()
    main_window.show()

    # sync wordgen options
    wordgen.add_headings = config.get("add_headings", True)
    wordgen.add_page_breaks = config.get("add_page_breaks", False)

    try:
        sys.exit(app.exec())
    except Exception:
        logger.exception("Qt app exit error")

if __name__ == "__main__":
    try:
        logger.info("Starting ScreenshotAssist - full single file (with active window metadata)")
        main()
    except Exception:
        logger.exception("Fatal error in main")
        print("An error occurred - check", LOG_FILE)