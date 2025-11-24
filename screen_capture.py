import pyautogui
from docx import Document
from datetime import datetime
import os
import keyboard

# ========== CONFIGURATION ==========
DOC_PATH = "Screenshot_Report.docx"  # Output Word file
HOTKEY = "f9"                        # Key to trigger screenshot
# ==================================

# Create or open document
if os.path.exists(DOC_PATH):
    doc = Document(DOC_PATH)
else:
    doc = Document()
    doc.add_heading("Automated Screenshot Report", 0)
    doc.add_paragraph("Press F9 anytime to capture a screenshot.\n")

print("✅ Screenshot Agent is running...")
print("👉 Press F9 to take a screenshot, or press ESC to exit.\n")

while True:
    try:
        # If hotkey is pressed
        if keyboard.is_pressed(HOTKEY):
            timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            filename = f"screenshot_{timestamp}.png"
            pyautogui.screenshot(filename)

            doc.add_heading(f"Screenshot taken at {timestamp}", level=1)
            doc.add_picture(filename, width=None)
            doc.save(DOC_PATH)

            print(f"📸 Captured and added to {DOC_PATH}")
            # Wait until key is released (avoid multiple triggers)
            keyboard.wait('f9', suppress=True)

        # Exit if ESC pressed
        if keyboard.is_pressed("esc"):
            print("🛑 Agent stopped.")
            break

    except KeyboardInterrupt:
        print("Exiting...")
        break