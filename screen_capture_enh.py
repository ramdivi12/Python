import tkinter as tk
from tkinter import ttk
import pyautogui
from docx import Document
from datetime import datetime
import threading
import time
import keyboard
import os

# ===================== CONFIG ======================
DOC_PATH = "Screenshot_Report.docx"
HOTKEY = "f9"
running = False
scheduled_running = False
# ===================================================

# Create/open Word document
if os.path.exists(DOC_PATH):
    doc = Document(DOC_PATH)
else:
    doc = Document()
    doc.add_heading("Automated Screenshot Report", 0)

lock = threading.Lock()

def take_screenshot():
    """Capture screenshot and add to Word document."""
    with lock:
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"screenshot_{timestamp}.png"
        pyautogui.screenshot(filename)

        doc.add_heading(f"Screenshot - {timestamp}", level=1)
        doc.add_picture(filename)
        doc.save(DOC_PATH)

def hotkey_listener():
    """Listen for F9 key while running."""
    global running
    while running:
        if keyboard.is_pressed(HOTKEY):
            take_screenshot()
            time.sleep(1)

def scheduled_capture(interval):
    """Take screenshots automatically every X seconds."""
    global scheduled_running
    while scheduled_running:
        take_screenshot()
        time.sleep(interval)

def start_manual():
    """Start manual screenshot mode."""
    global running
    running = True
    threading.Thread(target=hotkey_listener, daemon=True).start()

def stop_manual():
    """Stop manual capture."""
    global running
    running = False

def start_schedule():
    """Start scheduled screenshot mode."""
    global scheduled_running
    try:
        sec = int(schedule_entry.get())
    except:
        status_var.set("❌ Invalid interval")
        return
    scheduled_running = True
    threading.Thread(target=scheduled_capture, args=(sec,), daemon=True).start()
    status_var.set(f"📸 Scheduled every {sec} seconds")

def stop_schedule():
    global scheduled_running
    scheduled_running = False
    status_var.set("⛔ Schedule stopped")

# ------------------- UI --------------------

window = tk.Tk()
window.title("Screenshot Capture Agent")
window.geometry("400x280")

status_var = tk.StringVar()
status_var.set("Ready...")

lbl_manual = ttk.Label(window, text="Manual Capture (Press F9)", font=("Arial", 12))
lbl_manual.pack(pady=10)

btn_start = ttk.Button(window, text="Start Manual Mode", command=start_manual)
btn_start.pack()

btn_stop = ttk.Button(window, text="Stop Manual Mode", command=stop_manual)
btn_stop.pack()

ttk.Separator(window).pack(fill="x", pady=10)

lbl_schedule = ttk.Label(window, text="Scheduled Capture", font=("Arial", 12))
lbl_schedule.pack()

schedule_entry = ttk.Entry(window)
schedule_entry.insert(0, "5")  # default 5 sec
schedule_entry.pack()

btn_start_schedule = ttk.Button(window, text="Start Schedule", command=start_schedule)
btn_start_schedule.pack(pady=5)

btn_stop_schedule = ttk.Button(window, text="Stop Schedule", command=stop_schedule)
btn_stop_schedule.pack()

lbl_status = ttk.Label(window, textvariable=status_var, font=("Arial", 10))
lbl_status.pack(pady=15)

window.mainloop()